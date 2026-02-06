from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from pymongo import MongoClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import PyPDF2
import docx
import io
import pyphen

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
MONGO_URL = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
DB_NAME = os.environ.get('DB_NAME', 'flowread_db')
client = MongoClient(MONGO_URL)
db = client[DB_NAME]
documents_table = db['documents']
sessions_table = db['reading_sessions']

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Initialize pyphen for syllable detection
dic = pyphen.Pyphen(lang='en')

# Models
class Document(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    content: str
    word_count: int
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    file_type: str

class DocumentCreate(BaseModel):
    title: str
    content: str
    word_count: int
    file_type: str

class ReadingSession(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    current_word_index: int
    total_words: int
    words_read: int
    time_spent: int  # in seconds
    speed_wpm: int
    last_updated: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    completed: bool = False

class ReadingSessionCreate(BaseModel):
    document_id: str
    current_word_index: int
    total_words: int
    words_read: int
    time_spent: int
    speed_wpm: int
    completed: bool = False

class ReadingSessionUpdate(BaseModel):
    current_word_index: Optional[int] = None
    words_read: Optional[int] = None
    time_spent: Optional[int] = None
    speed_wpm: Optional[int] = None
    completed: Optional[bool] = None

class UserStats(BaseModel):
    total_documents: int
    total_words_read: int
    total_time_spent: int
    average_speed: int
    documents_completed: int

class ProcessedWord(BaseModel):
    word: str
    syllables: List[str]
    vowels: List[int]  # indices of vowels in each syllable

# Helper functions
def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    try:
        docx_file = io.BytesIO(file_content)
        doc = docx.Document(docx_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")

def process_text_to_words(text: str) -> List[str]:
    # Clean and split text into words
    words = text.split()
    # Remove empty strings and clean words
    words = [w.strip() for w in words if w.strip()]
    return words

def detect_syllables(word: str) -> List[str]:
    # Use pyphen to detect syllables
    syllables = dic.inserted(word).split('-')
    return syllables if len(syllables) > 1 else [word]

def is_vowel(char: str) -> bool:
    return char.lower() in 'aeiou'

def get_vowel_indices(syllable: str) -> List[int]:
    return [i for i, char in enumerate(syllable) if is_vowel(char)]

# Routes
@api_router.get("/")
async def root():
    return {"message": "FlowRead API"}

@api_router.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    try:
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        filename = file.filename.lower()
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(content)
            file_type = 'pdf'
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(content)
            file_type = 'docx'
        elif filename.endswith('.txt'):
            text = content.decode('utf-8')
            file_type = 'txt'
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF, DOCX, or TXT.")
        
        # Process text
        words = process_text_to_words(text)
        word_count = len(words)
        
        if word_count == 0:
            raise HTTPException(status_code=400, detail="No text found in the document.")
        
        # Create document
        doc_data = DocumentCreate(
            title=file.filename,
            content=text,
            word_count=word_count,
            file_type=file_type
        )
        
        doc_obj = Document(**doc_data.model_dump())
        doc_dict = doc_obj.model_dump()
        
        documents_table.insert_one(doc_dict)
        
        return doc_obj
    
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Upload error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@api_router.get("/documents", response_model=List[Document])
async def get_documents():
    documents = list(documents_table.find().sort('created_at', -1).limit(100))
    for doc in documents:
        doc['id'] = doc.get('id') or str(doc.get('_id'))
        if '_id' in doc: del doc['_id']
    return documents

@api_router.get("/documents/{doc_id}", response_model=Document)
async def get_document(doc_id: str):
    doc = documents_table.find_one({"id": doc_id})
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if '_id' in doc: del doc['_id']
    return doc

@api_router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    result = documents_table.delete_one({"id": doc_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Also delete associated sessions
    sessions_table.delete_many({"document_id": doc_id})
    
    return {"message": "Document deleted successfully"}

@api_router.get("/documents/{doc_id}/words")
async def get_document_words(doc_id: str):
    doc = documents_table.find_one({"id": doc_id})
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    words = process_text_to_words(doc['content'])
    
    # Process words with syllables
    processed_words = []
    for word in words:
        syllables = detect_syllables(word)
        vowel_data = [get_vowel_indices(syl) for syl in syllables]
        processed_words.append({
            "word": word,
            "syllables": syllables,
            "vowels": vowel_data
        })
    
    return {"words": processed_words}

@api_router.post("/sessions", response_model=ReadingSession)
async def create_session(session_data: ReadingSessionCreate):
    session_obj = ReadingSession(**session_data.model_dump())
    session_dict = session_obj.model_dump()
    
    sessions_table.insert_one(session_dict)
    
    return session_obj

@api_router.put("/sessions/{session_id}", response_model=ReadingSession)
async def update_session(session_id: str, update_data: ReadingSessionUpdate):
    session = sessions_table.find_one({"id": session_id})
    
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Update fields
    update_dict = {k: v for k, v in update_data.model_dump().items() if v is not None}
    update_dict['last_updated'] = datetime.now(timezone.utc).isoformat()
    
    sessions_table.update_one({"id": session_id}, {"$set": update_dict})
    
    # Get updated session
    updated_session = sessions_table.find_one({"id": session_id})
    if '_id' in updated_session: del updated_session['_id']
    
    return ReadingSession(**updated_session)

@api_router.get("/sessions/document/{doc_id}", response_model=ReadingSession)
async def get_latest_session(doc_id: str):
    session = sessions_table.find_one(
        {"document_id": doc_id},
        sort=[("last_updated", -1)]
    )
    
    if not session:
        return None
    
    if '_id' in session: del session['_id']
    return session

@api_router.get("/stats", response_model=UserStats)
async def get_stats():
    total_documents = documents_table.count_documents({})
    
    sessions = list(sessions_table.find())
    
    total_words_read = sum(s['words_read'] for s in sessions)
    total_time_spent = sum(s['time_spent'] for s in sessions)
    documents_completed = len([s for s in sessions if s['completed']])
    
    # Calculate average speed
    speeds = [s['speed_wpm'] for s in sessions if s['speed_wpm'] > 0]
    average_speed = int(sum(speeds) / len(speeds)) if speeds else 0
    
    return UserStats(
        total_documents=total_documents,
        total_words_read=total_words_read,
        total_time_spent=total_time_spent,
        average_speed=average_speed,
        documents_completed=documents_completed
    )

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)